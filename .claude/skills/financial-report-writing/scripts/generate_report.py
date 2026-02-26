"""
金融研究报告格式化生成器
========================

纯输出层：接收预处理好的数据和分析结论，生成符合 GB/T 9704-2012 的 Word 文档。
不包含任何数据获取或分析逻辑。

Usage:
    from generate_report import ReportGenerator

    gen = ReportGenerator()
    gen.generate_report(
        title="三一重工投资研究报告",
        sections=[
            {"heading": "核心提要", "points": ["ROE 15.2%，盈利能力较强", ...]},
            {"heading": "一、投资评级", "table": {...}},
            {"heading": "二、核心观点", "paragraphs": [...]},
            {"heading": "三、市场表现回顾", "paragraphs": [...], "table": {...}},
        ],
        output_path="tmp/session/output/report.docx"
    )
"""

import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from document_formatter import GBT9704Formatter


class ReportGenerator:
    """
    通用金融研究报告生成器

    职责：接收结构化内容，输出格式化的 Word 文档。
    不关心数据从哪里来、分析怎么做。
    """

    def __init__(self):
        self.doc = None
        self.formatter = None

    def generate_report(
        self,
        title: str,
        sections: List[Dict],
        organization: str = "",
        date_str: str = None,
        include_cover: bool = False,
        include_disclaimer: bool = False,
        data_source: str = "Tushare",
        output_path: str = None,
    ) -> str:
        """
        生成格式化的 Word 研究报告

        Args:
            title: 报告标题，如 "三一重工投资研究报告"
            sections: 报告章节列表，每个章节是一个 dict：
                - heading: str — 章节标题
                - level: int — 标题级别（1=一级，2=二级），默认 1
                - points: List[str] — 编号要点（用于核心提要等）
                - paragraphs: List[str] — 正文段落
                - table: dict — 表格数据
                    - headers: List[str]
                    - rows: List[List[str]]
                - bullets: List[str] — 项目符号列表
            organization: 研究机构名称（封面用）
            date_str: 日期字符串（封面用），默认今天
            include_cover: 是否包含封面页
            include_disclaimer: 是否包含免责声明
            data_source: 数据来源标注，默认 "Tushare"
            output_path: 输出路径

        Returns:
            保存的文件路径
        """
        # 初始化文档
        self.doc = Document()
        self.formatter = GBT9704Formatter(self.doc)

        if date_str is None:
            date_str = datetime.now().strftime("%Y年%m月%d日")

        # 封面
        if include_cover:
            self.formatter.add_cover_page(
                title=title,
                organization=organization,
                date_str=date_str,
            )

        # 免责声明
        if include_disclaimer:
            self._add_disclaimer()

        # 逐章节渲染
        for section in sections:
            self._render_section(section)

        # 数据来源注脚
        self.formatter.add_data_source_footnote(f"数据来源：{data_source}")

        # 页脚页码
        self.formatter.add_page_number_footer()

        # 保存
        if output_path is None:
            date_tag = datetime.now().strftime("%Y%m%d")
            safe_title = title.replace(" ", "_")[:30]
            output_path = f"{safe_title}_{date_tag}.docx"

        self.doc.save(output_path)
        return output_path

    def _render_section(self, section: Dict):
        """渲染单个章节"""
        heading = section.get("heading", "")
        level = section.get("level", 1)

        if heading:
            self.formatter.add_heading(heading, level=level)

        # 编号要点
        if "points" in section:
            for i, point in enumerate(section["points"], 1):
                p = self.doc.add_paragraph()
                run = p.add_run(f"{i}. {point}")
                self.formatter.set_font(run, "仿宋", 14)
                from docx.enum.text import WD_LINE_SPACING
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Pt(28)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)

        # 正文段落
        if "paragraphs" in section:
            for text in section["paragraphs"]:
                self.formatter.add_paragraph(text)

        # 表格
        if "table" in section:
            tbl = section["table"]
            self.formatter.add_three_line_table(
                tbl["rows"],
                tbl["headers"],
                font_name="仿宋",
                size=12,
            )

        # 项目符号
        if "bullets" in section:
            for item in section["bullets"]:
                p = self.doc.add_paragraph()
                run = p.add_run(f"• {item}")
                self.formatter.set_font(run, "仿宋", 14)
                from docx.enum.text import WD_LINE_SPACING
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Pt(28)

    def _add_disclaimer(self):
        """添加免责声明"""
        self.formatter.add_heading("免责声明", level=1)
        self.formatter.add_paragraph(
            "本报告所载资料的来源及观点皆为公开信息，但并不能保证其准确性和完整性。"
            "本报告仅供参考，不构成任何投资建议或承诺，投资者应审慎决策，独立判断，"
            "自行承担投资风险。"
        )


if __name__ == "__main__":
    # 快速测试：生成一个空壳报告验证格式
    gen = ReportGenerator()
    path = gen.generate_report(
        title="测试公司投资研究报告",
        organization="测试证券研究所",
        sections=[
            {
                "heading": "核心提要",
                "points": [
                    "公司 ROE 为 15.2%，盈利能力较强",
                    "当前 PE(TTM) 25.3 倍，估值合理",
                    "技术面震荡上行，均线多头排列",
                    '给予"增持"评级，目标价 35 元',
                ],
            },
            {
                "heading": "一、投资评级",
                "table": {
                    "headers": ["投资建议", "目标价格", "有效期", "风险等级"],
                    "rows": [["增持", "35.00元", "2027年2月26日", "中等风险"]],
                },
            },
            {
                "heading": "二、核心观点",
                "paragraphs": [
                    "公司基本面健康，盈利能力处于行业前列。",
                    "技术面呈震荡上行态势，MA5 > MA10 > MA20，均线多头排列。",
                ],
            },
        ],
        output_path="/tmp/test_report.docx",
    )
    print(f"测试报告已保存: {path}")
