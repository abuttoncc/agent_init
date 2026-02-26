"""
Word文档格式化模块
按照GB/T 9704-2012标准格式化金融研究报告
"""

from typing import List, Tuple

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class GBT9704Formatter:
    """
    GB/T 9704-2012党政机关公文格式化器
    """

    # 字体映射（号数 -> 磅值）
    FONT_SIZES = {
        '三号': 16,
        '四号': 14,
        '小四': 12,
        '五号': 10.5
    }

    def __init__(self, doc: Document):
        """
        初始化格式化器

        Args:
            doc: Document对象
        """
        self.doc = doc
        self._setup_page()

    def _setup_page(self):
        """设置页面格式"""
        for section in self.doc.sections:
            # 页面边距（单位：毫米）
            section.top_margin = Mm(37)
            section.bottom_margin = Mm(35)
            section.left_margin = Mm(28)
            section.right_margin = Mm(26)

    def set_font(self, run, font_name: str, size: int = None, bold: bool = None, color: RGBColor = None):
        """
        设置字体

        Args:
            run: Run对象
            font_name: 字体名称（中文字体）
            size: 字号（磅值），如 16, 14, 12
            bold: 是否加粗
            color: 字体颜色，默认黑色
        """
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        # 默认黑色，确保标题等不会继承 Word 内建样式颜色
        run.font.color.rgb = color if color is not None else RGBColor(0, 0, 0)

    def add_heading(self, text: str, level: int = 1, font_name: str = '黑体', size: int = 16):
        """
        添加标题

        Args:
            text: 标题文本
            level: 标题级别（1-3）
            font_name: 字体名称
            size: 字号

        Returns:
            Paragraph对象
        """
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            run = heading.runs[0]
            self.set_font(run, font_name, size, bold=True)
            heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(12)
            return heading
        else:
            # 二级、三级标题使用"一、"、"(一)"等格式
            prefix_map = {1: '一、', 2: '（一）', 3: '1.'}
            prefix = prefix_map.get(level, '')

            paragraph = self.doc.add_paragraph(f"{prefix}{text}")
            run = paragraph.runs[0]
            self.set_font(run, font_name, size, bold=True)

            # 段落格式
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)

            return paragraph

    def add_paragraph(self, text: str,
                     font_name: str = '仿宋',
                     size: int = 14,
                     first_line_indent: bool = True,
                     alignment: str = 'justify'):
        """
        添加正文段落

        Args:
            text: 段落文本
            font_name: 字体名称
            size: 字号（磅值）
            first_line_indent: 是否首行缩进
            alignment: 对齐方式（'left', 'center', 'right', 'justify'）

        Returns:
            Paragraph对象
        """
        paragraph = self.doc.add_paragraph(text)
        run = paragraph.runs[0]
        self.set_font(run, font_name, size)

        # 段落格式
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.paragraph_format.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        if first_line_indent:
            # 首行缩进2字符（28pt）
            paragraph.paragraph_format.first_line_indent = Pt(28)

        # 1.5倍行距
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 段间距1行
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(12)

        return paragraph

    def add_bullet_list(self, items: List[str], font_name: str = '仿宋', size: int = 14):
        """
        添加项目符号列表

        Args:
            items: 列表项
            font_name: 字体名称
            size: 字号
        """
        for item in items:
            paragraph = self.doc.add_paragraph(item, style='List Bullet')
            run = paragraph.runs[0]
            self.set_font(run, font_name, size)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)

    def add_three_line_table(self, data: List[List[str]],
                            headers: List[str],
                            font_name: str = '仿宋',
                            size: int = 12):
        """
        添加三线表

        Args:
            data: 表格数据（二维列表）
            headers: 表头
            font_name: 字体名称
            size: 字号

        Returns:
            Table对象
        """
        # 创建表格
        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        self._set_three_line_table_borders(table)

        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            self._set_cell_font(cell, header, font_name, size, bold=True)
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置数据行
        for row_idx, row_data in enumerate(data, start=1):
            cells = table.rows[row_idx].cells
            for i, cell_data in enumerate(row_data):
                cell = cells[i]
                self._set_cell_font(cell, str(cell_data), font_name, size)
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return table

    def _set_three_line_table_borders(self, table):
        """设置三线表边框"""
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')

        # 顶底边框 1.5pt (sz=18)
        for pos in ['top', 'bottom']:
            border = OxmlElement(f'w:{pos}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '18')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)

        # 内部横线 0.75pt (sz=9)
        for pos in ['insideH']:
            border = OxmlElement(f'w:{pos}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '9')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)

        # 无左右和竖线
        for pos in ['left', 'right', 'insideV']:
            border = OxmlElement(f'w:{pos}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _set_cell_font(self, cell, text: str, font_name: str, size: int, bold: bool = False):
        """设置单元格字体"""
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        self.set_font(run, font_name, size, bold=bold)

    def add_page_number_footer(self):
        """添加页码页脚"""
        for section in self.doc.sections:
            footer = section.footer

            # 添加页码段落
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 页码格式：— 1 —
            self.add_footer_run(p, '— ', font_name='仿宋', size=12)

            # 添加页码字段
            self.add_page_number_field(p)

            self.add_footer_run(p, ' ', font_name='仿宋', size=12)
            self.add_footer_run(p, '—', font_name='仿宋', size=12)

    def add_footer_run(self, paragraph, text: str, font_name: str, size: int):
        """添加页脚run"""
        run = paragraph.add_run(text)
        self.set_font(run, font_name, size)

    def add_page_number_field(self, paragraph):
        """添加页码字段"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run = paragraph.add_run()
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        self.set_font(run, '仿宋', 12)

    def add_cover_page(self, title: str, organization: str, date_str: str):
        """
        添加封面页

        Args:
            title: 报告标题
            organization: 机构名称
            date_str: 日期
        """
        # 标题
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(title)
        self.set_font(run, '黑体', 22, bold=True)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(100)
        paragraph.paragraph_format.space_after = Pt(50)

        # 机构名称
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(organization)
        self.set_font(run, '黑体', 16)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(20)

        # 日期
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(date_str)
        self.set_font(run, '黑体', 16)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加分页符
        self.doc.add_page_break()

    def add_data_source_footnote(self, text: str = "数据来源：Tushare"):
        """
        添加数据来源注脚

        Args:
            text: 注脚文本
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        self.set_font(run, '仿宋', 12)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# 便捷函数
def create_formatted_document() -> Tuple[Document, GBT9704Formatter]:
    """
    创建格式化的Word文档

    Returns:
        (Document, GBT9704Formatter) 文档对象和格式化器
    """
    doc = Document()
    formatter = GBT9704Formatter(doc)
    return doc, formatter


if __name__ == '__main__':
    # 测试代码
    doc, formatter = create_formatted_document()

    # 添加封面
    formatter.add_cover_page(
        '某某股票投资研究报告',
        '某某证券研究所',
        '2026年2月6日'
    )

    # 添加核心提要
    formatter.add_heading('核心提要', level=1)
    formatter.add_bullet_list([
        '该公司2025年三季度净资产收益率（ROE）为15.2%，盈利能力较强',
        '当前市盈率（TTM）为25.3倍，略高于行业平均水平',
        '技术面显示股价呈现震荡上行趋势，建议关注回调机会',
        '给予"买入"评级，目标价格35元'
    ])

    # 添加投资评级表格
    formatter.add_heading('一、投资评级', level=1)
    data = [
        ['买入', '35.00元', '2026年12月31日', '中风险']
    ]
    formatter.add_three_line_table(
        data,
        ['投资建议', '目标价格', '有效期', '风险等级']
    )

    # 添加数据来源
    formatter.add_data_source_footnote()

    # 保存文档
    output_path = 'test_report.docx'
    doc.save(output_path)
    print(f"测试报告已保存到: {output_path}")
